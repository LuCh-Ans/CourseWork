import pdfplumber
import PyPDF2
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer, LTChar, LTRect, LTFigure
from PIL import Image
from pdf2image import convert_from_path
import pytesseract 
import os
from pathlib import Path
from docx import Document

from dotenv import load_dotenv

load_dotenv()
 
pytesseract.pytesseract.tesseract_cmd = os.getenv("TESSERACT_PATH", r"tesseract")
POPPLER_PATH = None

def extract_text(file_path: Path) -> list[dict]:
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return extract_text_docx(file_path)
    elif suffix == ".pdf":
        return extract_text_pdf(file_path)
    elif suffix == ".txt":
        raw = file_path.read_text(encoding="utf-8", errors="ignore")
        return [{
            "content": line.strip(),
            "metadata": {
                "source": file_path.name,
                "section": "General",
                "type": "text"
            }
        } for line in raw.split('\n') if line.strip()]
    else:
        raise ValueError(f"Формат {suffix} не поддерживается процессором.")

def extract_text_pdf(pdf: Path) -> list[dict]:
    content_blocks: list[dict] = []
    pdf_path_str = str(pdf)
    curr_section = 'General'
    images = convert_from_path(pdf_path_str, poppler_path=POPPLER_PATH)
    with pdfplumber.open(pdf_path_str) as f:
        for page_num, page in enumerate(f.pages):
            page_text = page.extract_text()
            if not page_text or len(page_text.strip()) < 10:
                image = images[page_num]
                page_text = pytesseract.image_to_string(image, lang='rus+eng')
                block_type = "ocr"
            else:
                block_type = "text"
            if page_text:
                for line in page_text.split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                    is_heading = len(line) < 80 and not line.endswith('.')
            
                    if is_heading:
                        curr_section = line
                        content_blocks.append({
                            "content": line,
                            "metadata": {
                                "source": pdf.name,
                                "section": curr_section,
                                "type": "heading",
                                "page": page_num + 1
                            }
                        })
                    else:
                        content_blocks.append({
                            "content": line,
                            "metadata": {
                                "source": pdf.name,
                                "section": curr_section,
                                "type": "text",
                                "page": page_num + 1
                            }
                        })
            tables = page.extract_tables()
            for table_idx, table in enumerate(tables):
                for row_idx, row in enumerate(table):
                    row_text = " | ".join(map(str, filter(None, row)))
                    if row_text.strip():
                        content_blocks.append({
                            "content": row_text,
                            "metadata": {
                                "source": pdf.name,
                                "section": curr_section,
                                "type": "table",
                                "page": page_num + 1,
                                "table_index": table_idx,
                                "row_index": row_idx
                            }
                        })
     
     
    return content_blocks

def extract_text_docx(file_path: Path) -> list[dict]:
    doc = Document(file_path)
    content_blocks: list[dict] = []
    curr_section = "General"
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name and paragraph.style.name.startswith('Heading'):
            curr_section = text
            content_blocks.append({
                "content": text,
                "metadata": {
                    "source": file_path.name,
                    "section": curr_section,
                    "type": "heading",
                    "level": paragraph.style.name   
                }
            })
            continue
        content_blocks.append({
            "content": text,
            "metadata": {
                "source": file_path.name,
                "section": curr_section,
                "type": "text"
            }
        })
    for table_in, table in enumerate(doc.tables):
        for row_in, row in enumerate(table.rows):
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                content_blocks.append({
                    "content": " | ".join(row_text),
                    "metadata": {
                        "source": file_path.name,
                        "section": curr_section,
                        "type": "table",
                        "table_index": table_in,
                        "row_index": row_in
                    }
                })
    
    return content_blocks

def clean_text(text: str) -> str:
    line = text.split('\n')
    result_text = []
    for i in line:
        i = i.strip()
        if i:
            result_text.append(i)
    result = " ".join(result_text)
    return result



def process_file(file_path: Path) -> list[dict]:
    blocks = extract_text(file_path)
    blocks = [
        {**b, "content": clean_text(b["content"])}
        for b in blocks
        if clean_text(b["content"])
    ]
    return blocks

def blocks_to_text(blocks: list[dict], max_length: int = 120_000) -> str:
    lines = []
    total_length = 0
    truncated = False

    for block in blocks:
        meta = block["metadata"]
        content = block["content"]

        if meta["type"] == "heading":
            line = f"\n## {content}\n"
        elif meta["type"] == "table":
            line = f"\n[Таблица — раздел «{meta['section']}»]\n{content}\n"
        elif meta["type"] == "ocr":
            line = f"\n[OCR, стр. {meta.get('page', '?')}]\n{content}\n"
        else:
            line = f"\n{content}"

        if total_length + len(line) > max_length:
            truncated = True
            break

        lines.append(line)
        total_length += len(line)

    text = "\n".join(lines).strip()
    if truncated:
        text += "\n\n... (документ обрезан)"
    return text

def create_chunks(blocks: list, chunk_size: int = 600, chunk_overlap: int = 150):
    """
    Разбивает текст блоков на чанки с перекрытием.
    chunk_size: макс. количество символов в одном фрагменте.
    chunk_overlap: сколько символов из конца предыдущего чанка попадет в начало следующего.
    """
    full_text = " ".join([b["content"] for b in blocks])
    chunks = []
    if len(full_text) <= chunk_size:
        return [{"content": full_text}]

    start = 0
    while start < len(full_text):
        end = start + chunk_size
        chunk_content = full_text[start:end]
        chunks.append({"content": chunk_content.strip()})
        start += (chunk_size - chunk_overlap)
    return chunks
 